"""
webWhys - Website & Competitor Analysis Tool
Copyright (C) 2026 Nicole Scott

This program is free software: you can redistribute it and/or modify
it under the terms of the GNU Affero General Public License as published
by the Free Software Foundation, either version 3 of the License, or
(at your option) any later version.

This program is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
GNU Affero General Public License for more details.

You should have received a copy of the GNU Affero General Public License
along with this program. If not, see <https://www.gnu.org/licenses/>.

---

Website Competitor Scanner - Main FastAPI Application
Analyzes your website, competitors, and brand documents to provide
SEO/GEO/LLM discoverability optimization suggestions.
"""

import os
import io
import tempfile
from typing import Optional
from datetime import datetime
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import uvicorn

from scraper import WebsiteScraper
from analyzer import OptimizationAnalyzer
from document_processor import DocumentProcessor
from metric_explanations import generate_metric_insights, get_all_explanations
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_bookmark(paragraph, bookmark_name):
    """Add a bookmark to a paragraph for TOC linking."""
    run = paragraph.add_run()
    bookmarkStart = OxmlElement('w:bookmarkStart')
    bookmarkStart.set(qn('w:id'), '0')
    bookmarkStart.set(qn('w:name'), bookmark_name)
    bookmarkEnd = OxmlElement('w:bookmarkEnd')
    bookmarkEnd.set(qn('w:id'), '0')
    run._element.addprevious(bookmarkStart)
    run._element.addnext(bookmarkEnd)

def add_hyperlink(paragraph, text, bookmark_name):
    """Add a hyperlink to a bookmark within the document."""
    # Create hyperlink element pointing to the bookmark
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)

    # Create run with formatting
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add blue color and underline for link appearance
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')  # Blue
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    # Add text
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)

def add_external_hyperlink(paragraph, part, text, url, color_rgb='C9C3BD'):
    """Add a clickable external hyperlink to a paragraph."""
    # part should be either doc.part (main document) or footer/header part
    # Add relationship to document
    rel_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), rel_id)

    # Create run with formatting
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add color and underline for link appearance
    color = OxmlElement('w:color')
    color.set(qn('w:val'), color_rgb)
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    # Add text
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)

def shade_table_header(table, color_rgb):
    """Add shading to the header row of a table."""
    for cell in table.rows[0].cells:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color_rgb)
        cell._element.get_or_add_tcPr().append(shading_elm)
        # Make header text bold and white
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

def format_metrics_table(table):
    """Format metric tables to be narrower and properly spaced."""
    # Set column widths - make first column wider for labels
    if len(table.columns) == 2:
        table.columns[0].width = Inches(2.5)
        table.columns[1].width = Inches(2.0)
    # Center align the table
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # Set table alignment to left with margins
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '4500')
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

def shade_section(paragraph, color_hex):
    """Add background shading to a paragraph section."""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)

def set_table_borders(table, color_hex='auto', size='single'):
    """Set table border color and style."""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Create table borders element
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')  # Border size
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color_hex)
        tblBorders.append(border)

    tblPr.append(tblBorders)

def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Create table borders element with no border
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)

    tblPr.append(tblBorders)

def add_divider_line(doc, color_hex='D3D3D3'):
    """Add a horizontal divider line."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

def indent_paragraph(paragraph, indent_inches=0.0625):
    """Add left indentation to a paragraph."""
    paragraph.paragraph_format.left_indent = Inches(indent_inches)

app = FastAPI(
    title="Website Competitor Scanner",
    description="Scan your website and competitors to get AI-powered optimization suggestions",
    version="1.0.0"
)

# CORS middleware for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files
app.mount("/static", StaticFiles(directory="static"), name="static")


class ScanRequest(BaseModel):
    your_website: str
    competitor_urls: list[str]
    focus_areas: Optional[list[str]] = None


class AnalysisResponse(BaseModel):
    status: str
    your_site_analysis: dict
    competitor_analyses: list[dict]
    recommendations: list[dict]
    copy_suggestions: list[dict] = []
    priority_actions: list[dict]


@app.get("/", response_class=HTMLResponse)
async def serve_frontend():
    """Serve the main frontend application."""
    return FileResponse("static/index.html")


@app.post("/api/scan", response_model=AnalysisResponse)
async def scan_websites(
    your_website: str = Form(...),
    competitor_urls: str = Form(default=""),  # Comma-separated, optional
    focus_areas: Optional[str] = Form(None),
    brand_docs: list[UploadFile] = File(default=[])
):
    """
    Scan your website and optionally competitor websites, analyze uploaded documents,
    and return prioritized optimization suggestions.
    """
    try:
        scraper = WebsiteScraper()
        analyzer = OptimizationAnalyzer()
        doc_processor = DocumentProcessor()

        # Parse competitor URLs (optional)
        competitors = [url.strip() for url in competitor_urls.split(",") if url.strip()] if competitor_urls else []

        # Parse focus areas if provided
        areas = []
        if focus_areas:
            areas = [area.strip() for area in focus_areas.split(",") if area.strip()]

        # Scrape your website
        your_site_data = await scraper.analyze_website(your_website)

        # Scrape competitor websites
        competitor_data = []
        for comp_url in competitors[:5]:  # Limit to 5 competitors
            try:
                comp_analysis = await scraper.analyze_website(comp_url)
                competitor_data.append(comp_analysis)
            except Exception as e:
                competitor_data.append({
                    "url": comp_url,
                    "error": str(e),
                    "status": "failed"
                })

        # Process uploaded documents (temporary, no storage)
        brand_context = []
        for doc in brand_docs:
            if doc.filename:
                # Create temp file, process, then delete
                with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(doc.filename)[1]) as tmp:
                    content = await doc.read()
                    tmp.write(content)
                    tmp_path = tmp.name

                try:
                    doc_content = doc_processor.extract_content(tmp_path, doc.filename)
                    doc_content["filename"] = doc.filename
                    brand_context.append(doc_content)
                finally:
                    # Clean up temp file
                    os.unlink(tmp_path)

        # Generate optimization recommendations
        recommendations = await analyzer.generate_recommendations(
            your_site=your_site_data,
            competitors=competitor_data,
            brand_documents=brand_context,
            focus_areas=areas
        )

        # Generate metric insights with explanations
        metric_insights = generate_metric_insights(your_site_data, competitor_data)

        return {
            "status": "success",
            "your_site_analysis": your_site_data,
            "competitor_analyses": competitor_data,
            "recommendations": recommendations["recommendations"],
            "copy_suggestions": recommendations.get("copy_suggestions", []),
            "priority_actions": recommendations["priority_actions"],
            "metric_insights": metric_insights,
            "metric_explanations": get_all_explanations()
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/quick-scan")
async def quick_scan(request: ScanRequest):
    """Quick scan without file uploads - JSON API."""
    scraper = WebsiteScraper()
    analyzer = OptimizationAnalyzer()

    your_site_data = await scraper.analyze_website(request.your_website)

    competitor_data = []
    for comp_url in request.competitor_urls[:5]:
        try:
            comp_analysis = await scraper.analyze_website(comp_url)
            competitor_data.append(comp_analysis)
        except Exception as e:
            competitor_data.append({
                "url": comp_url,
                "error": str(e),
                "status": "failed"
            })

    recommendations = await analyzer.generate_recommendations(
        your_site=your_site_data,
        competitors=competitor_data,
        brand_documents=[],
        focus_areas=request.focus_areas or []
    )

    return {
        "status": "success",
        "your_site_analysis": your_site_data,
        "competitor_analyses": competitor_data,
        "recommendations": recommendations["recommendations"],
        "priority_actions": recommendations["priority_actions"]
    }


@app.get("/api/health")
async def health_check():
    """Health check endpoint."""
    return {"status": "healthy", "service": "Website Competitor Scanner"}


class ExportRequest(BaseModel):
    """Request model for docx export."""
    your_site_analysis: dict
    competitor_analyses: list[dict]
    recommendations: list[dict]
    priority_actions: list[dict]
    copy_suggestions: Optional[list[dict]] = None
    metric_insights: Optional[list[dict]] = None


@app.post("/api/export-docx")
async def export_docx(data: ExportRequest):
    """
    Generate a formatted Word document report that opens in Google Docs.
    """
    doc = Document()

    # Define styles with improved spacing and formatting
    title_style = doc.styles['Title']
    title_style.font.color.rgb = RGBColor(0x4A, 0x34, 0x53)  # Your brand headline color (deep plum)
    title_style.font.name = 'Avenir'
    title_style.font.size = Pt(28)
    title_style.paragraph_format.space_after = Pt(6)

    heading1_style = doc.styles['Heading 1']
    heading1_style.font.color.rgb = RGBColor(0x4A, 0x34, 0x53)  # Same as title
    heading1_style.font.name = 'Avenir'
    heading1_style.font.size = Pt(16)
    heading1_style.font.bold = True
    heading1_style.paragraph_format.space_before = Pt(12)
    heading1_style.paragraph_format.space_after = Pt(8)

    heading2_style = doc.styles['Heading 2']
    heading2_style.font.color.rgb = RGBColor(0xB3, 0x8B, 0x91)  # Your brand highlight-3 (mauve)
    heading2_style.font.name = 'Avenir'
    heading2_style.font.size = Pt(13)
    heading2_style.font.bold = True
    heading2_style.paragraph_format.space_before = Pt(10)
    heading2_style.paragraph_format.space_after = Pt(6)
    heading2_style.paragraph_format.left_indent = Inches(0)

    heading3_style = doc.styles['Heading 3']
    heading3_style.font.color.rgb = RGBColor(0x4A, 0x34, 0x53)
    heading3_style.font.name = 'Avenir'
    heading3_style.font.size = Pt(11)
    heading3_style.font.bold = True
    heading3_style.paragraph_format.space_before = Pt(8)
    heading3_style.paragraph_format.space_after = Pt(4)

    # Set default body font to Avenir with better spacing
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Avenir'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)

    # Title
    doc.add_heading('webWhys Analysis', 0)

    # Date and URL
    analysis = data.your_site_analysis
    seo = analysis.get("seo_factors", {})

    date_para = doc.add_paragraph()
    date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}").italic = True
    date_para.add_run(f"\nURL Analyzed: {analysis.get('url', 'N/A')}")

    # Extra spacing before TOC/TLDR
    doc.add_paragraph()
    doc.add_paragraph()

    # TOC and TL;DR side-by-side layout
    layout_table = doc.add_table(rows=1, cols=2)
    layout_table.style = 'Table Grid'
    layout_table.columns[0].width = Inches(2.25)
    layout_table.columns[1].width = Inches(2.75)
    remove_table_borders(layout_table)  # Remove borders from layout table

    # Left cell: Table of Contents
    left_cell = layout_table.rows[0].cells[0]
    left_para = left_cell.paragraphs[0]
    left_heading = left_para.add_run('Table of Contents')
    left_heading.bold = True
    left_heading.font.size = Pt(13)
    left_heading.font.color.rgb = RGBColor(0x4A, 0x34, 0x53)  # Brand color

    toc_items = [
        ('Top Priority Actions', 'TopPriorityActions'),
        ('Site Analysis Metrics', 'SiteAnalysisMetrics'),
        ('All Recommendations', 'AllRecommendations'),
        ('Suggested Wording', 'SuggestedWording'),
        ('Competitor Comparison', 'CompetitorComparison')
    ]

    # Add Key Insights if available
    if data.metric_insights:
        toc_items.insert(2, ('Key Insights vs Competitors', 'KeyInsights'))

    for item_name, bookmark_id in toc_items:
        p = left_cell.add_paragraph(style='List Bullet')
        add_hyperlink(p, item_name, bookmark_id)
        p_format = p.paragraph_format
        p_format.left_indent = Inches(0.125)
        p_format.space_before = Pt(0)
        p_format.space_after = Pt(2)
        # Fix hyperlink color and formatting
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)  # Blue hyperlink color

    # Right cell: TL;DR
    right_cell = layout_table.rows[0].cells[1]
    right_para = right_cell.paragraphs[0]
    right_heading = right_para.add_run('TL;DR')
    right_heading.bold = True
    right_heading.font.size = Pt(13)
    right_heading.font.color.rgb = RGBColor(0x4A, 0x34, 0x53)  # Brand color

    tldr_para = right_cell.add_paragraph(
        "This analysis evaluates your website's visibility across search engines and AI-powered "
        "platforms. Below are prioritized recommendations to improve your rankings, answer-engine "
        "discoverability, and brand positioning clarity."
    )
    tldr_para.paragraph_format.left_indent = Inches(0.0625)  # 1/16" indent

    # Extra spacing before Priority Actions section
    doc.add_paragraph()
    doc.add_paragraph()
    add_divider_line(doc, 'E8E6E3')  # Pale beige divider

    # Priority Actions
    heading = doc.add_heading('Top Priority Actions', level=1)
    add_bookmark(heading.insert_paragraph_before(''), 'TopPriorityActions')
    # Don't indent headings - only body text

    for i, action in enumerate(data.priority_actions or [], 1):
        # Action title with number
        p = doc.add_paragraph()
        p.add_run(f"{i}. {action.get('title', 'Untitled')}").bold = True
        p.paragraph_format.space_after = Pt(0)  # No space after title
        p.paragraph_format.left_indent = Inches(0.0625)  # Indent

        # Metadata on second line
        meta = doc.add_paragraph()
        meta.add_run(f"Category: {action.get('category', 'N/A')} | Impact: {action.get('impact', 'N/A')} | Effort: {action.get('effort', 'N/A')}")
        meta.paragraph_format.space_after = Pt(0)  # No space after metadata
        meta.paragraph_format.left_indent = Inches(0.25)  # Indented (1/4")

    # Page break before Site Analysis Metrics
    doc.add_page_break()

    # Site Analysis Metrics
    heading = doc.add_heading('Site Analysis Metrics', level=1)
    add_bookmark(heading.insert_paragraph_before(''), 'SiteAnalysisMetrics')

    # SEO Factors table
    seo_heading = doc.add_heading('SEO Factors', level=2)

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'

    # Set column widths: 33% / 64%
    table.columns[0].width = Inches(1.54)  # 33%
    table.columns[1].width = Inches(2.98)  # 64%

    # Set table width to 75% of page (approximately 4.88 inches for standard 6.5" usable width)
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Set table indent to 0.5 inches (720 twips)
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '720')  # 0.5 inches in twips
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)

    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '7020')  # 75% of page width in twips (4.88 inches)
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # Set medium grey borders
    set_table_borders(table, '999999')

    seo_rows = [
        ('Title', seo.get('title', 'Missing')[:60] + '...' if len(seo.get('title', '')) > 60 else seo.get('title', 'Missing')),
        ('Title Length', f"{seo.get('title_length', 0)} chars (optimal: 50-60)"),
        ('Meta Description Length', f"{seo.get('meta_description_length', 0)} chars (optimal: 150-160)"),
        ('H1 Tags', f"{len(seo.get('h1_tags', []))} (optimal: 1)"),
        ('Word Count', str(seo.get('word_count', 0))),
        ('Images Missing Alt Text', str(seo.get('images_without_alt', 0)))
    ]

    for i, (label, value) in enumerate(seo_rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
        # Add light shading to all cells
        for cell in table.rows[i].cells:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F6F3F1')
            cell._element.get_or_add_tcPr().append(shading_elm)

    doc.add_paragraph()  # Spacing after table
    add_divider_line(doc, 'E8E6E3')  # Pale beige divider

    # Technical Factors & LLM Discoverability (side-by-side)
    layout_table = doc.add_table(rows=1, cols=2)
    layout_table.style = 'Table Grid'
    remove_table_borders(layout_table)  # Remove borders from layout table
    # Set column widths for layout table
    layout_table.columns[0].width = Inches(2.5)
    layout_table.columns[1].width = Inches(2.5)

    # Left cell: Technical Factors
    left_cell = layout_table.rows[0].cells[0]
    left_para = left_cell.paragraphs[0]
    left_heading = left_para.add_run('Technical Factors')
    left_heading.bold = True
    left_heading.font.size = Pt(13)
    left_heading.font.color.rgb = RGBColor(0xB3, 0x8B, 0x91)
    left_para.paragraph_format.left_indent = Inches(0.0625)  # 1/16" indent

    tech = analysis.get("technical_factors", {})
    tech_table = left_cell.add_table(rows=4, cols=2)
    tech_table.style = 'Table Grid'
    remove_table_borders(tech_table)  # No borders
    format_metrics_table(tech_table)

    tech_rows = [
        ('HTTPS', 'Yes' if tech.get('https') else 'No'),
        ('Sitemap', 'Yes' if tech.get('has_sitemap') else 'No'),
        ('Robots.txt', 'Yes' if tech.get('has_robots_txt') else 'No'),
        ('Mobile Viewport', 'Yes' if tech.get('mobile_friendly_hints') else 'No')
    ]

    for i, (label, value) in enumerate(tech_rows):
        tech_table.rows[i].cells[0].text = label
        tech_table.rows[i].cells[1].text = value

    # Right cell: LLM Discoverability
    right_cell = layout_table.rows[0].cells[1]
    right_para = right_cell.paragraphs[0]
    right_heading = right_para.add_run('LLM Discoverability')
    right_heading.bold = True
    right_heading.font.size = Pt(13)
    right_heading.font.color.rgb = RGBColor(0xB3, 0x8B, 0x91)
    right_para.paragraph_format.left_indent = Inches(0.0625)  # 1/16" indent

    llm = analysis.get("llm_discoverability", {})
    llm_table = right_cell.add_table(rows=4, cols=2)
    llm_table.style = 'Table Grid'
    remove_table_borders(llm_table)  # No borders
    format_metrics_table(llm_table)

    llm_rows = [
        ('Structured Content', 'Yes' if llm.get('structured_content') else 'No'),
        ('FAQ Schema', 'Yes' if llm.get('faq_schema') else 'No'),
        ('How-To Schema', 'Yes' if llm.get('how_to_schema') else 'No'),
        ('External Citations', str(llm.get('citations_and_sources', 0)))
    ]

    for i, (label, value) in enumerate(llm_rows):
        llm_table.rows[i].cells[0].text = label
        llm_table.rows[i].cells[1].text = value

    doc.add_paragraph()  # Spacing after layout table
    add_divider_line(doc, 'E8E6E3')  # Pale beige divider

    # GEO Factors & Scannability & Messaging (side-by-side)
    layout_table = doc.add_table(rows=1, cols=2)
    layout_table.style = 'Table Grid'
    remove_table_borders(layout_table)  # Remove borders from layout table
    # Set column widths for layout table
    layout_table.columns[0].width = Inches(2.5)
    layout_table.columns[1].width = Inches(2.5)

    # Left cell: GEO Factors
    left_cell = layout_table.rows[0].cells[0]
    left_para = left_cell.paragraphs[0]
    left_heading = left_para.add_run('GEO (AI Citation) Factors')
    left_heading.bold = True
    left_heading.font.size = Pt(13)
    left_heading.font.color.rgb = RGBColor(0xB3, 0x8B, 0x91)
    left_para.paragraph_format.left_indent = Inches(0.0625)  # 1/16" indent

    geo = analysis.get("geo_factors", {})
    geo_table = left_cell.add_table(rows=4, cols=2)
    geo_table.style = 'Table Grid'
    remove_table_borders(geo_table)  # No borders
    format_metrics_table(geo_table)

    geo_rows = [
        ('Citation Ready', 'Yes' if geo.get('citation_ready') else 'No'),
        ('Statistics Present', 'Yes' if geo.get('statistics_present') else 'No'),
        ('Comparison Tables', 'Yes' if geo.get('comparison_tables') else 'No'),
        ('Lists/Bullet Points', str(geo.get('lists_and_bullets', 0)))
    ]

    for i, (label, value) in enumerate(geo_rows):
        geo_table.rows[i].cells[0].text = label
        geo_table.rows[i].cells[1].text = value

    # Right cell: Scannability & Messaging
    right_cell = layout_table.rows[0].cells[1]
    right_para = right_cell.paragraphs[0]
    right_heading = right_para.add_run('Scannability & Messaging')
    right_heading.bold = True
    right_heading.font.size = Pt(13)
    right_heading.font.color.rgb = RGBColor(0xB3, 0x8B, 0x91)
    right_para.paragraph_format.left_indent = Inches(0.0625)  # 1/16" indent

    scannability = analysis.get("scannability", {})
    messaging = analysis.get("page_messaging", {})

    # Remove Word Count and Content Organization - now 5 rows instead of 7
    scan_table = right_cell.add_table(rows=5, cols=2)
    scan_table.style = 'Table Grid'
    remove_table_borders(scan_table)  # No borders
    format_metrics_table(scan_table)

    scannability_rows = [
        ('Heading Count', str(scannability.get('heading_count', 0))),
        ('Heading Hierarchy', scannability.get('heading_hierarchy_quality', 'Unknown')),
        ('Lists/Bullet Points', str(scannability.get('list_count', 0))),
        ('Avg Paragraph', f"{scannability.get('avg_paragraph_length', 0)} words"),
        ('Primary Message', 'Clear' if scannability.get('has_clear_primary_message') else 'Unclear')
    ]

    for i, (label, value) in enumerate(scannability_rows):
        scan_table.rows[i].cells[0].text = label
        scan_table.rows[i].cells[1].text = str(value)

    doc.add_paragraph()  # Spacing after layout table
    add_divider_line(doc, 'E8E6E3')  # Pale beige divider

    # Metric Insights (if available)
    if data.metric_insights:
        heading = doc.add_heading('Key Insights vs Competitors', level=1)
        add_bookmark(heading.insert_paragraph_before(''), 'KeyInsights')
        for insight in data.metric_insights:
            p = doc.add_paragraph()
            p.add_run(f"{insight.get('metric', 'Unknown').replace('_', ' ').title()}").bold = True
            p.add_run(f" - {insight.get('status', 'unknown').upper()}")

            value_para = doc.add_paragraph(f"Your value: {insight.get('your_value', 'N/A')} | Competitor average: {insight.get('competitor_avg', 'N/A')}")
            value_para.paragraph_format.left_indent = Inches(0.0625)  # 1/16" indent for body text

            if insight.get('explanation'):
                exp_para = doc.add_paragraph()
                exp_para.add_run("Why it matters: ").italic = True
                exp_para.add_run(insight['explanation'])
                exp_para.paragraph_format.left_indent = Inches(0.0625)  # 1/16" indent for body text

            if insight.get('recommendation'):
                rec_para = doc.add_paragraph()
                rec_para.add_run("Recommendation: ").bold = True
                rec_para.add_run(insight['recommendation'])
                rec_para.paragraph_format.left_indent = Inches(0.0625)  # 1/16" indent for body text

            doc.add_paragraph()

    # Divider before All Recommendations
    add_divider_line(doc, 'E8E6E3')

    # All Recommendations
    heading = doc.add_heading('All Recommendations', level=1)
    add_bookmark(heading.insert_paragraph_before(''), 'AllRecommendations')

    for i, rec in enumerate(data.recommendations or [], 1):
        # Recommendation title
        p = doc.add_paragraph()
        p.add_run(f"{i}. {rec.get('title', 'Untitled')}").bold = True

        # Metadata on second line
        meta = doc.add_paragraph()
        meta.add_run(f"Category: {rec.get('category', 'General')} | Impact: {rec.get('impact', 'N/A')} | Effort: {rec.get('effort', 'N/A')}")
        meta.paragraph_format.left_indent = Inches(0.25)  # Indented (1/4")

        # Description
        if rec.get('description'):
            desc_para = doc.add_paragraph(rec['description'])
            desc_para.paragraph_format.left_indent = Inches(0.0625)  # 1/16" indent for body text

        # Action steps
        if rec.get('specific_actions'):
            doc.add_paragraph("Action Steps:").runs[0].bold = True
            for step in rec['specific_actions']:
                step_para = doc.add_paragraph(style='List Bullet')
                step_run = step_para.add_run(step)
                step_run.font.size = Pt(9.5)  # 9.5pt font for text (not bullet)
                step_para.paragraph_format.left_indent = Inches(0.5)  # 1/2" indent for bullets
                step_para.paragraph_format.first_line_indent = Inches(-0.15)  # hanging indent for tighter spacing

        # Expected outcome
        if rec.get('expected_outcome'):
            outcome_para = doc.add_paragraph()
            outcome_para.add_run("Expected Outcome: ").bold = True
            outcome_para.add_run(rec['expected_outcome'])
            outcome_para.paragraph_format.left_indent = Inches(0.0625)  # 1/16" indent for body text

        doc.add_paragraph()  # Spacing

    # Suggested Wording (Copy Suggestions)
    if data.copy_suggestions and len(data.copy_suggestions) > 0:
        heading = doc.add_heading('Suggested Wording', level=1)
        add_bookmark(heading.insert_paragraph_before(''), 'SuggestedWording')

        for suggestion in data.copy_suggestions:
            category = suggestion.get('category', 'Unknown')
            current = suggestion.get('current', '')
            why = suggestion.get('why', '')
            suggestions = suggestion.get('suggestions', [])

            # Category heading
            p = doc.add_paragraph()
            p.add_run(category).bold = True

            # Current text
            if current:
                current_para = doc.add_paragraph(style='List Bullet')
                current_run = current_para.add_run(f"Current: {current}")
                current_run.font.size = Pt(9.5)  # 9.5pt font for text (not bullet)
                current_para.paragraph_format.left_indent = Inches(0.5)  # 1/2" indent for bullets
                current_para.paragraph_format.first_line_indent = Inches(-0.15)  # hanging indent for tighter spacing

            # Why explanation
            if why:
                why_para = doc.add_paragraph(style='List Bullet')
                why_run = why_para.add_run(f"Why: {why}")
                why_run.font.size = Pt(9.5)  # 9.5pt font for text (not bullet)
                why_para.paragraph_format.left_indent = Inches(0.5)  # 1/2" indent for bullets
                why_para.paragraph_format.first_line_indent = Inches(-0.15)  # hanging indent for tighter spacing

            # Suggested options
            if suggestions:
                doc.add_paragraph("Suggested options:").runs[0].italic = True
                for suggestion_text in suggestions:
                    sugg_para = doc.add_paragraph(style='List Bullet')
                    sugg_run = sugg_para.add_run(suggestion_text)
                    sugg_run.font.size = Pt(9.5)  # 9.5pt font for text (not bullet)
                    sugg_para.paragraph_format.left_indent = Inches(0.5)  # 1/2" indent for bullets
                    sugg_para.paragraph_format.first_line_indent = Inches(-0.15)  # hanging indent for tighter spacing

            doc.add_paragraph()  # Spacing between suggestions

    # Page break before Competitor Comparison section
    doc.add_page_break()

    # Divider before Competitor Comparison
    add_divider_line(doc, 'E8E6E3')

    # Competitor Comparison
    if data.competitor_analyses:
        heading = doc.add_heading('Competitor Comparison', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center the heading
        add_bookmark(heading.insert_paragraph_before(''), 'CompetitorComparison')

        # Comprehensive metrics covering SEO, Technical, LLM, GEO, and Scannability
        metrics = [
            # SEO Metrics
            ('Word Count', lambda a: str(a.get('seo_factors', {}).get('word_count', 'N/A'))),
            ('Title Length', lambda a: f"{a.get('seo_factors', {}).get('title_length', 0)} chars"),
            ('H1 Tags', lambda a: str(len(a.get('seo_factors', {}).get('h1_tags', [])))),
            ('Meta Description', lambda a: '✓' if a.get('seo_factors', {}).get('meta_description') else '✗'),

            # Technical Metrics
            ('HTTPS', lambda a: '✓' if a.get('technical_factors', {}).get('https') else '✗'),
            ('Mobile Friendly', lambda a: '✓' if a.get('technical_factors', {}).get('mobile_friendly_hints') else '✗'),
            ('Sitemap', lambda a: '✓' if a.get('technical_factors', {}).get('has_sitemap') else '✗'),

            # Scannability & Messaging
            ('Heading Count', lambda a: str(a.get('scannability', {}).get('heading_count', 0))),
            ('Lists Present', lambda a: str(a.get('scannability', {}).get('list_count', 0))),

            # LLM Discoverability
            ('Structured Data', lambda a: '✓' if a.get('content_analysis', {}).get('has_structured_data') else '✗'),
            ('FAQ Schema', lambda a: '✓' if a.get('llm_discoverability', {}).get('faq_schema') else '✗'),
            ('How-To Schema', lambda a: '✓' if a.get('llm_discoverability', {}).get('how_to_schema') else '✗'),

            # GEO Factors (AI Citation)
            ('Statistics Present', lambda a: '✓' if a.get('geo_factors', {}).get('statistics_present') else '✗'),
            ('Comparison Tables', lambda a: '✓' if a.get('geo_factors', {}).get('comparison_tables') else '✗'),
            ('Citation Ready', lambda a: '✓' if a.get('geo_factors', {}).get('citation_ready') else '✗'),
            ('Issues Found', lambda a: str(len(a.get('issues', []))))
        ]

        # Create comparison table with expanded metrics
        num_comps = len(data.competitor_analyses)
        num_metrics = len(metrics)
        table = doc.add_table(rows=num_metrics + 1, cols=2 + num_comps)
        table.style = 'Table Grid'

        # Header row
        table.rows[0].cells[0].text = 'Metric'
        table.rows[0].cells[1].text = 'Your Site'
        for i, comp in enumerate(data.competitor_analyses):
            table.rows[0].cells[2 + i].text = comp.get('domain', comp.get('url', f'Competitor {i+1}'))[:20]

        for row_idx, (metric_name, get_value) in enumerate(metrics, 1):
            table.rows[row_idx].cells[0].text = metric_name
            table.rows[row_idx].cells[1].text = get_value(analysis)
            for comp_idx, comp in enumerate(data.competitor_analyses):
                table.rows[row_idx].cells[2 + comp_idx].text = get_value(comp)

        # Shade the header row with brand color (deep plum)
        shade_table_header(table, '4A3453')

        doc.add_paragraph()  # Spacing after table

        # Messaging Breakdown section
        doc.add_heading('Messaging Breakdown', level=2)
        msg_intro = doc.add_paragraph(
            "What each site is trying to say, who they're talking to, and the value prop a visitor would walk away with."
        )
        msg_intro.paragraph_format.left_indent = Inches(0.0625)  # 1/16" indent for body text

        # Light beige-grey shading color
        subtle_shade = 'F5F3F1'

        # Your site messaging
        your_site_heading = doc.add_heading('Your Site', level=3)
        your_site_heading.paragraph_format.left_indent = Inches(0.125)  # 1/8" indent
        your_messaging = analysis.get("page_messaging", {})
        messaging_data = [
            ("Primary Headline (H1)", your_messaging.get('primary_message', 'N/A')),
            ("Apparent Audience", your_messaging.get('apparent_audience', 'N/A')),
            ("Value Prop Visitor Would Walk Away With", your_messaging.get('value_proposition', 'N/A')),
            ("Tone", your_messaging.get('tone', 'N/A')),
            ("CTA Language", "; ".join(your_messaging.get('cta_language', [])) if your_messaging.get('cta_language') else 'N/A'),
            ("Keywords They Appear to Be Targeting", "; ".join(your_messaging.get('keyword_targets', [])) if your_messaging.get('keyword_targets') else 'N/A')
        ]

        msg_table = doc.add_table(rows=len(messaging_data), cols=2)
        msg_table.style = 'Table Grid'

        # Set table properties
        msg_tbl = msg_table._element
        msg_tblPr = msg_tbl.tblPr
        if msg_tblPr is None:
            msg_tblPr = OxmlElement('w:tblPr')
            msg_tbl.insert(0, msg_tblPr)

        # Set table indent to match SEO table (0.5 inches = 720 twips)
        msg_tblInd = OxmlElement('w:tblInd')
        msg_tblInd.set(qn('w:w'), '720')
        msg_tblInd.set(qn('w:type'), 'dxa')
        msg_tblPr.append(msg_tblInd)

        # Set table to fixed layout to respect column widths
        msg_tblLayout = OxmlElement('w:tblLayout')
        msg_tblLayout.set(qn('w:type'), 'fixed')
        msg_tblPr.append(msg_tblLayout)

        msg_tblW = OxmlElement('w:tblW')
        msg_tblW.set(qn('w:w'), '8000')  # Table width in twips (5.556 inches)
        msg_tblW.set(qn('w:type'), 'dxa')
        msg_tblPr.append(msg_tblW)

        # Populate table and set cell widths for proper 33%/67% split
        for i, (label, value) in enumerate(messaging_data):
            msg_table.rows[i].cells[0].text = label
            msg_table.rows[i].cells[1].text = value if value else 'N/A'

            # Set cell width for proper column proportions (33% / 67%)
            cell0_tcPr = msg_table.rows[i].cells[0]._element.get_or_add_tcPr()
            cell0_tcW = OxmlElement('w:tcW')
            cell0_tcW.set(qn('w:w'), '2640')  # 33% of 8000 twips
            cell0_tcW.set(qn('w:type'), 'dxa')
            # Remove any existing tcW before appending
            existing_tcW = cell0_tcPr.find(qn('w:tcW'))
            if existing_tcW is not None:
                cell0_tcPr.remove(existing_tcW)
            cell0_tcPr.append(cell0_tcW)

            # Set width for column 2
            cell1_tcPr = msg_table.rows[i].cells[1]._element.get_or_add_tcPr()
            cell1_tcW = OxmlElement('w:tcW')
            cell1_tcW.set(qn('w:w'), '5360')  # 67% of 8000 twips
            cell1_tcW.set(qn('w:type'), 'dxa')
            existing_tcW = cell1_tcPr.find(qn('w:tcW'))
            if existing_tcW is not None:
                cell1_tcPr.remove(existing_tcW)
            cell1_tcPr.append(cell1_tcW)

            # Shade the label column
            shade_cell = OxmlElement('w:shd')
            shade_cell.set(qn('w:fill'), subtle_shade)
            msg_table.rows[i].cells[0]._element.get_or_add_tcPr().append(shade_cell)
            # Make label bold
            for paragraph in msg_table.rows[i].cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        doc.add_paragraph()  # Spacing after table

        # Competitor messaging
        for comp_idx, comp in enumerate(data.competitor_analyses):
            comp_url = comp.get('domain', comp.get('url', f'Competitor {comp_idx+1}'))
            comp_heading = doc.add_heading(comp_url, level=3)
            comp_heading.paragraph_format.left_indent = Inches(0.125)  # 1/8" indent

            comp_messaging = comp.get("page_messaging", {})
            messaging_data = [
                ("Primary Headline (H1)", comp_messaging.get('primary_message', 'N/A')),
                ("Apparent Audience", comp_messaging.get('apparent_audience', 'N/A')),
                ("Value Prop Visitor Would Walk Away With", comp_messaging.get('value_proposition', 'N/A')),
                ("Tone", comp_messaging.get('tone', 'N/A')),
                ("CTA Language", "; ".join(comp_messaging.get('cta_language', [])) if comp_messaging.get('cta_language') else 'N/A'),
                ("Keywords They Appear to Be Targeting", "; ".join(comp_messaging.get('keyword_targets', [])) if comp_messaging.get('keyword_targets') else 'N/A')
            ]

            msg_table = doc.add_table(rows=len(messaging_data), cols=2)
            msg_table.style = 'Table Grid'

            # Set table properties
            msg_tbl = msg_table._element
            msg_tblPr = msg_tbl.tblPr
            if msg_tblPr is None:
                msg_tblPr = OxmlElement('w:tblPr')
                msg_tbl.insert(0, msg_tblPr)

            # Set table indent to match SEO table (0.5 inches = 720 twips)
            msg_tblInd = OxmlElement('w:tblInd')
            msg_tblInd.set(qn('w:w'), '720')
            msg_tblInd.set(qn('w:type'), 'dxa')
            msg_tblPr.append(msg_tblInd)

            # Set table to fixed layout to respect column widths
            msg_tblLayout = OxmlElement('w:tblLayout')
            msg_tblLayout.set(qn('w:type'), 'fixed')
            msg_tblPr.append(msg_tblLayout)

            msg_tblW = OxmlElement('w:tblW')
            msg_tblW.set(qn('w:w'), '8000')  # Table width in twips (5.556 inches)
            msg_tblW.set(qn('w:type'), 'dxa')
            msg_tblPr.append(msg_tblW)

            # Populate table and set cell widths for proper 33%/67% split
            for i, (label, value) in enumerate(messaging_data):
                msg_table.rows[i].cells[0].text = label
                msg_table.rows[i].cells[1].text = value if value else 'N/A'

                # Set cell width for proper column proportions (33% / 67%)
                cell0_tcPr = msg_table.rows[i].cells[0]._element.get_or_add_tcPr()
                cell0_tcW = OxmlElement('w:tcW')
                cell0_tcW.set(qn('w:w'), '2640')  # 33% of 8000 twips
                cell0_tcW.set(qn('w:type'), 'dxa')
                # Remove any existing tcW before appending
                existing_tcW = cell0_tcPr.find(qn('w:tcW'))
                if existing_tcW is not None:
                    cell0_tcPr.remove(existing_tcW)
                cell0_tcPr.append(cell0_tcW)

                # Set width for column 2
                cell1_tcPr = msg_table.rows[i].cells[1]._element.get_or_add_tcPr()
                cell1_tcW = OxmlElement('w:tcW')
                cell1_tcW.set(qn('w:w'), '5360')  # 67% of 8000 twips
                cell1_tcW.set(qn('w:type'), 'dxa')
                existing_tcW = cell1_tcPr.find(qn('w:tcW'))
                if existing_tcW is not None:
                    cell1_tcPr.remove(existing_tcW)
                cell1_tcPr.append(cell1_tcW)

                # Shade the label column
                shade_cell = OxmlElement('w:shd')
                shade_cell.set(qn('w:fill'), subtle_shade)
                msg_table.rows[i].cells[0]._element.get_or_add_tcPr().append(shade_cell)
                # Make label bold
                for paragraph in msg_table.rows[i].cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            doc.add_paragraph()  # Spacing after table

    # Divider before footer
    add_divider_line(doc, 'E8E6E3')

    # Add footer to document sections with clickable external hyperlinks
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Get the footer part for hyperlink relationships
    footer_part = footer.part

    # Add webWhys as clickable hyperlink to GitHub
    add_external_hyperlink(footer_para, footer_part, 'webWhys', 'https://github.com/nicole-os/mvpmm', 'C9C3BD')

    # Add separator
    sep_run = footer_para.add_run(' | ')
    sep_run.italic = True
    sep_run.font.size = Pt(8)
    sep_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Add MVP Marketing as clickable hyperlink to LinkedIn
    add_external_hyperlink(footer_para, footer_part, 'Minimum Viable Product (Marketing)', 'https://www.linkedin.com/in/nicolescottfromraleigh/', 'C9C3BD')

    footer_para.paragraph_format.space_before = Pt(6)

    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Generate filename
    filename = f"webWhys-report-{datetime.now().strftime('%d-%m-%y')}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
