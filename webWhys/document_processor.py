"""
webWhys - Website & Competitor Analysis Tool
Copyright (C) 2026 Nicole Scott

This program is free software: you can redistribute it and/or modify
it under the terms of the GNU Affero General Public License as published
by the Free Software Foundation, either version 3 of the License, or
(at your option) any later version.

This program is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
GNU Affero General Public License for more details.

You should have received a copy of the GNU Affero General Public License
along with this program. If not, see <https://www.gnu.org/licenses/>.

---

Document Processor Module
Extracts content from uploaded documents for brand/positioning analysis.
Supports PDF, DOCX, TXT, and Markdown files.
No long-term storage - processes in memory/temp files only.
"""

import os
import re
from typing import Optional


class DocumentProcessor:
    """Process uploaded documents to extract brand and positioning content."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".rtf"}
    MAX_CONTENT_LENGTH = 80000  # Characters

    def extract_content(self, file_path: str, original_filename: str) -> dict:
        """
        Extract content from a document file.
        Returns structured content for brand analysis.
        """
        ext = os.path.splitext(original_filename.lower())[1]

        if ext not in self.SUPPORTED_EXTENSIONS:
            return {
                "status": "unsupported",
                "error": f"Unsupported file type: {ext}",
                "content": "",
                "metadata": {}
            }

        try:
            if ext == ".pdf":
                content, metadata = self._extract_pdf(file_path)
            elif ext in {".docx", ".doc"}:
                content, metadata = self._extract_docx(file_path)
            elif ext == ".txt":
                content, metadata = self._extract_text(file_path)
            elif ext == ".md":
                content, metadata = self._extract_markdown(file_path)
            elif ext == ".rtf":
                content, metadata = self._extract_rtf(file_path)
            else:
                content, metadata = "", {}

            # Truncate if too long
            if len(content) > self.MAX_CONTENT_LENGTH:
                content = content[:self.MAX_CONTENT_LENGTH] + "\n\n[Content truncated...]"

            # Extract key brand elements
            brand_elements = self._extract_brand_elements(content)

            return {
                "status": "success",
                "content": content,
                "metadata": metadata,
                "brand_elements": brand_elements,
                "word_count": len(content.split()),
                "char_count": len(content)
            }

        except Exception as e:
            return {
                "status": "error",
                "error": str(e),
                "content": "",
                "metadata": {}
            }

    def _fix_pdf_word_joins(self, text: str) -> str:
        """
        Fix common PDF extraction artifacts where words on adjacent lines
        get concatenated without spaces (e.g. "workinstant" → "work instant").
        Strategy: insert a space before a lowercase letter that immediately
        follows a lowercase letter when the transition looks like a line-break fusion.
        We use a conservative regex — only split at camelCase-like boundaries
        inside what looks like a multi-word run.
        """
        import re
        # Insert space between a lowercase letter followed immediately by an
        # uppercase letter (catches "workInstant" style fusions from PDFs)
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
        # Collapse multiple spaces
        text = re.sub(r' {2,}', ' ', text)
        return text

    def _extract_pdf(self, file_path: str) -> tuple[str, dict]:
        """Extract text from PDF files."""
        try:
            import pypdf
            content = []
            metadata = {}

            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                metadata["pages"] = len(reader.pages)

                if reader.metadata:
                    metadata["title"] = reader.metadata.get("/Title", "")
                    metadata["author"] = reader.metadata.get("/Author", "")

                for page in reader.pages:
                    # Try layout mode first (pypdf ≥ 3.1) — preserves spacing better
                    try:
                        text = page.extract_text(extraction_mode="layout")
                    except TypeError:
                        text = page.extract_text()
                    if text:
                        content.append(self._fix_pdf_word_joins(text))

            return "\n\n".join(content), metadata
        except ImportError:
            # Fallback if pypdf not installed
            return self._extract_text(file_path)

    def _extract_docx(self, file_path: str) -> tuple[str, dict]:
        """Extract text from DOCX files."""
        try:
            import docx
            doc = docx.Document(file_path)
            content = []
            metadata = {
                "paragraphs": len(doc.paragraphs)
            }

            # Extract core properties if available
            if doc.core_properties:
                metadata["title"] = doc.core_properties.title or ""
                metadata["author"] = doc.core_properties.author or ""

            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        content.append(row_text)

            return "\n\n".join(content), metadata
        except ImportError:
            return "", {"error": "python-docx not installed"}

    def _extract_text(self, file_path: str) -> tuple[str, dict]:
        """Extract content from plain text files."""
        encodings = ["utf-8", "latin-1", "cp1252"]

        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    content = f.read()
                return content, {"encoding": encoding}
            except UnicodeDecodeError:
                continue

        return "", {"error": "Could not decode file"}

    def _extract_markdown(self, file_path: str) -> tuple[str, dict]:
        """Extract content from Markdown files."""
        content, metadata = self._extract_text(file_path)

        # Extract headers for metadata
        headers = re.findall(r"^#+\s+(.+)$", content, re.MULTILINE)
        metadata["headers"] = headers[:10]  # First 10 headers
        metadata["format"] = "markdown"

        return content, metadata

    def _extract_rtf(self, file_path: str) -> tuple[str, dict]:
        """Extract content from RTF files (basic extraction)."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                rtf_content = f.read()

            # Basic RTF stripping
            text = re.sub(r"\\[a-z]+\d*\s?", "", rtf_content)
            text = re.sub(r"[{}]", "", text)
            text = text.strip()

            return text, {"format": "rtf"}
        except Exception:
            return "", {"error": "Could not parse RTF"}

    def _extract_brand_elements(self, content: str) -> dict:
        """Extract brand-related elements from document content."""
        elements = {
            "mission_vision": [],
            "value_propositions": [],
            "key_differentiators": [],
            "target_audience": [],
            "tone_voice": [],
            "keywords": [],
            "competitive_positioning": []
        }

        content_lower = content.lower()

        # Look for mission/vision statements
        mission_patterns = [
            r"(?:our\s+)?mission[:\s]+([^.]+\.)",
            r"(?:our\s+)?vision[:\s]+([^.]+\.)",
            r"we\s+(?:are|help|enable|empower)\s+([^.]+\.)"
        ]
        for pattern in mission_patterns:
            matches = re.findall(pattern, content_lower, re.IGNORECASE)
            elements["mission_vision"].extend(matches[:3])

        # Look for value propositions
        value_patterns = [
            r"(?:we\s+)?(?:offer|provide|deliver)\s+([^.]+\.)",
            r"(?:our\s+)?(?:solution|product|service)\s+([^.]+\.)",
            r"benefits?\s*[:\-]\s*([^.]+\.)"
        ]
        for pattern in value_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            elements["value_propositions"].extend(matches[:5])

        # Look for differentiators
        diff_patterns = [
            r"(?:unlike|different from|compared to)\s+([^.]+\.)",
            r"(?:only|unique|first)\s+([^.]+\.)",
            r"what\s+(?:sets us apart|makes us different)\s*[:\-]?\s*([^.]+\.)"
        ]
        for pattern in diff_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            elements["key_differentiators"].extend(matches[:5])

        # Look for target audience mentions
        audience_patterns = [
            r"(?:for|designed for|built for|targeting)\s+([^.]+(?:teams?|companies|enterprises?|businesses?|organizations?)[^.]*\.)",
            r"(?:our\s+)?(?:customers?|clients?|users?)\s+(?:are|include)\s+([^.]+\.)"
        ]
        for pattern in audience_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            elements["target_audience"].extend(matches[:5])

        # Extract potential keywords (capitalized phrases, quoted terms)
        quoted = re.findall(r'"([^"]+)"', content)
        elements["keywords"].extend(quoted[:10])

        # Clean up and deduplicate
        for key in elements:
            elements[key] = list(set(elements[key]))[:5]

        return elements


class BrandContextBuilder:
    """Build a comprehensive brand context from multiple documents."""

    def build_context(self, documents: list[dict]) -> dict:
        """
        Combine multiple document analyses into a unified brand context.
        """
        context = {
            "combined_content": "",
            "all_brand_elements": {
                "mission_vision": [],
                "value_propositions": [],
                "key_differentiators": [],
                "target_audience": [],
                "keywords": []
            },
            "document_types": [],
            "total_word_count": 0
        }

        # Calculate per-doc character budget so all docs get fair representation
        successful_docs = [d for d in documents if d.get("status") == "success" and d.get("content")]
        per_doc_limit = 12000 // max(len(successful_docs), 1)

        for doc in successful_docs:
            # Combine content with per-doc limit so no single doc crowds out others
            context["combined_content"] += f"\n\n=== DOCUMENT: {doc.get('filename', 'Uploaded Document')} ===\n"
            context["combined_content"] += doc["content"][:per_doc_limit]
            context["total_word_count"] += doc.get("word_count", 0)

            # Merge brand elements
            if "brand_elements" in doc:
                for key, values in doc["brand_elements"].items():
                    if key in context["all_brand_elements"]:
                        context["all_brand_elements"][key].extend(values)

        # Deduplicate brand elements
        for key in context["all_brand_elements"]:
            context["all_brand_elements"][key] = list(set(context["all_brand_elements"][key]))[:10]

        return context
